#!/usr/bin/env python3

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer


def load_manifest(path_str: str):
    with open(path_str, "r", encoding="utf-8") as handle:
        return json.load(handle)


def build_pdf(manifest: dict, output_path: Path):
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ScrivixTitle",
        parent=styles["Title"],
        fontName="Times-Bold",
        fontSize=21,
        leading=26,
        spaceAfter=10,
    )
    heading_style = ParagraphStyle(
        "ScrivixHeading",
        parent=styles["Heading2"],
        fontName="Times-Bold",
        fontSize=14,
        leading=18,
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "ScrivixBody",
        parent=styles["BodyText"],
        fontName="Times-Roman",
        fontSize=11,
        leading=16,
        spaceAfter=8,
    )
    meta_style = ParagraphStyle(
        "ScrivixMeta",
        parent=styles["BodyText"],
        fontName="Times-Italic",
        fontSize=9,
        leading=13,
        textColor="#55636d",
        spaceAfter=6,
    )

    story = [
        Paragraph(manifest["title"], title_style),
        Paragraph(f'{manifest["subtitle"]} · {manifest["audience"]}', meta_style),
        Paragraph(f'Export profile: {manifest["profile_label"]}', meta_style),
        Spacer(1, 0.2 * inch),
    ]

    for section in manifest["sections"]:
        story.append(Paragraph(section["fileName"], heading_style))

        for paragraph in section["paragraphs"]:
            story.append(Paragraph(paragraph.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), body_style))

        story.append(Spacer(1, 0.12 * inch))

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.85 * inch,
        bottomMargin=0.85 * inch,
        title=manifest["title"],
    )
    doc.build(story)


def build_docx(manifest: dict, output_path: Path):
    document = Document()
    title = document.add_heading(manifest["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    meta = document.add_paragraph()
    meta.add_run(f'{manifest["subtitle"]} · {manifest["audience"]}').italic = True
    document.add_paragraph(f'Export profile: {manifest["profile_label"]}')

    for section in manifest["sections"]:
        document.add_heading(section["fileName"], level=1)

        for paragraph in section["paragraphs"]:
            document.add_paragraph(paragraph)

    document.save(str(output_path))


def main():
    if len(sys.argv) != 4:
        raise SystemExit("Usage: export_document.py <manifest.json> <format> <output>")

    manifest_path = sys.argv[1]
    output_format = sys.argv[2]
    output_path = Path(sys.argv[3])
    output_path.parent.mkdir(parents=True, exist_ok=True)
    manifest = load_manifest(manifest_path)

    if output_format == "pdf":
      build_pdf(manifest, output_path)
    elif output_format == "docx":
      build_docx(manifest, output_path)
    else:
      raise SystemExit(f"Unsupported output format: {output_format}")


if __name__ == "__main__":
    main()
